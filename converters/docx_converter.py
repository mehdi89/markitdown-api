import docx
import re
import logging

logger = logging.getLogger(__name__)

def convert_docx_to_markdown(file_path, use_llm=False):
    """
    Convert DOCX file to Markdown
    
    Args:
        file_path (str): Path to the DOCX file
        use_llm (bool): Whether to use LLM enhancement (not implemented)
        
    Returns:
        str: Markdown content
    """
    logger.debug(f"Converting DOCX file: {file_path}")
    
    try:
        doc = docx.Document(file_path)
        markdown_text = ""
        
        # Process document paragraphs
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
                
            # Check paragraph style for headings
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.replace('Heading', ''))
                markdown_text += f"{'#' * level} {para.text.strip()}\n\n"
            
            # Check for bullet lists
            elif para.style.name.startswith('List Bullet'):
                markdown_text += f"* {para.text.strip()}\n"
            
            # Check for numbered lists
            elif para.style.name.startswith('List Number'):
                # Extract number from text if possible
                match = re.match(r'^\s*(\d+)[\.)]?\s+(.+)$', para.text.strip())
                if match:
                    number, text = match.groups()
                    markdown_text += f"{number}. {text}\n"
                else:
                    markdown_text += f"1. {para.text.strip()}\n"
            
            # Regular paragraph
            else:
                # Check for bold and italic formatting
                text = para.text
                for run in para.runs:
                    if run.bold and run.italic:
                        text = text.replace(run.text, f"***{run.text}***")
                    elif run.bold:
                        text = text.replace(run.text, f"**{run.text}**")
                    elif run.italic:
                        text = text.replace(run.text, f"*{run.text}*")
                
                markdown_text += f"{text.strip()}\n\n"
        
        # Process tables
        for table in doc.tables:
            # Create table headers
            header_row = table.rows[0]
            headers = [cell.text.strip() for cell in header_row.cells]
            
            markdown_text += "| " + " | ".join(headers) + " |\n"
            markdown_text += "| " + " | ".join(["---"] * len(headers)) + " |\n"
            
            # Create table rows
            for i, row in enumerate(table.rows):
                if i == 0:  # Skip header row
                    continue
                cells = [cell.text.strip() for cell in row.cells]
                markdown_text += "| " + " | ".join(cells) + " |\n"
            
            markdown_text += "\n"
        
        # Clean up multiple newlines
        markdown_text = re.sub(r'\n{3,}', '\n\n', markdown_text)
        
        if use_llm:
            logger.info("LLM enhancement requested but not implemented for DOCX")
        
        return markdown_text
    
    except Exception as e:
        logger.error(f"Error converting DOCX to Markdown: {str(e)}")
        raise Exception(f"DOCX conversion error: {str(e)}")
